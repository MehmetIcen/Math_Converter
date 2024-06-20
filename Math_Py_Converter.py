import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import io
from sympy import symbols, sin, ln, sqrt, Rational, exp
from sympy.printing.latex import latex
from docx import Document
from docx.shared import Inches

#Make sure the install the requirments

# Create a new Document
doc = Document()
doc.add_heading('New_Math_Doc', 0)

# Define the symbol
x = symbols('x')

# List of similar functions in sympy format - How to type your functions @ ReadME
functions = [
    4*x / (1 + x**6)**Rational(1, 7),
    x**3 * sin(ln(x**3)),
    exp(x**Rational(1, 4)),
    3**(3*x),
    3*x / (1 + x**4)**Rational(1, 5),
    x**2 * sin(ln(x**2)),
    exp(sqrt(x)),
    4**(4*x)
]

# Helper function to create and save an image of a LaTeX formula
def create_latex_image(formula):
    plt.figure(figsize=(6, 2))
    latex_str = f'{latex(formula)}'
    plt.text(0.5, 0.5, f'${latex_str}$', fontsize=20, ha='center', va='center')
    plt.axis('off')
    buf = io.BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight')
    plt.close()
    buf.seek(0)
    return buf

# Add each function as an image to the document
for func in functions:
    buf = create_latex_image(func)
    doc.add_picture(buf, width=Inches(6.0))

# Save the document
file_path = r"C:\Users\MemoI\PycharmProjects\Math_Excerise_Converter\Your_Math_Exercise.docx"
doc.save(file_path)

#Returns the path of the new file
print(f"Document saved at {file_path}")
